import os
import re
import json
import logging
from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Setup logging to file and console
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log"),  # Save logs to a file
        logging.StreamHandler()          # Also log to console
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
DOCX_FILENAME = os.path.join(os.path.dirname(__file__), "generated_resume.docx")
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

@app.route('/')
def home():
    logger.info('Home route accessed.')
    return 'Welcome to the Resume DOCX Generator API!'

def replace_placeholder_with_format(paragraph, placeholder, value):
    """Replace placeholder while preserving formatting, including across multiple runs."""
    if placeholder not in paragraph.text:
        return

    # Combine all runs' text to find the placeholder
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return

    # Replace the placeholder in the combined text
    full_text = full_text.replace(placeholder, str(value))

    # Preserve the formatting of the first run
    first_run = paragraph.runs[0]
    formatting = {
        "bold": first_run.bold,
        "italic": first_run.italic,
        "underline": first_run.underline,
        "font_name": first_run.font.name,
        "font_size": first_run.font.size,
        "color": first_run.font.color.rgb,
    }

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Reassign the modified text to the first run and apply formatting
    first_run.text = full_text
    first_run.bold = formatting["bold"]
    first_run.italic = formatting["italic"]
    first_run.underline = formatting["underline"]
    if formatting["font_name"]:
        first_run.font.name = formatting["font_name"]
    if formatting["font_size"]:
        first_run.font.size = formatting["font_size"]
    if formatting["color"]:
        first_run.font.color.rgb = formatting["color"]

def replace_all_placeholders(doc, replacements):
    """Replace all placeholders in the document"""
    # Process all paragraphs in the document
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                replace_placeholder_with_format(paragraph, placeholder, value)
    
    # Process table cells if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            replace_placeholder_with_format(paragraph, placeholder, value)

    # Log placeholders that were not replaced
    for placeholder in replacements.keys():
        if any(placeholder in paragraph.text for paragraph in doc.paragraphs):
            logger.warning(f"Placeholder '{placeholder}' was not replaced in the document.")

def generate_docx(data):
    try:
        logger.info('Generating DOCX...')
        
        # Check if template exists
        if not os.path.exists(TEMPLATE_PATH):
            logger.error(f'Template not found at: {TEMPLATE_PATH}')
            raise FileNotFoundError(f"Template file not found: {TEMPLATE_PATH}")
        
        # Load the template
        doc = Document(TEMPLATE_PATH)
        logger.info(f'Template loaded from: {TEMPLATE_PATH}')
        
        # Debug: Print the content of the template
        logger.debug("Template content:")
        for paragraph in doc.paragraphs:
            logger.debug(paragraph.text)
        
        # Prepare replacements dictionary
        replacements = {
            '{{name}}': data.get('name', ''),
            '{{phone}}': data.get('phone', ''),
            '{{email}}': data.get('email', ''),
            '{{linkedin}}': data.get('linkedin', ''),
            '{{skills}}': ', '.join(data.get('skills', [])),
            '{{certification}}': data.get('certification', '')
        }
        
        # Handle experience section
        experiences = data.get('experience', [])
        for i, exp in enumerate(experiences):
            replacements[f'{{{{experience[{i}].designation}}}}'] = exp.get('designation', '')
            replacements[f'{{{{experience[{i}].dates}}}}'] = exp.get('dates', '')
            replacements[f'{{{{experience[{i}].company}}}}'] = exp.get('company', '')
            replacements[f'{{{{experience[{i}].city}}}}'] = exp.get('city', '')
            replacements[f'{{{{experience[{i}].1stbullet}}}}'] = exp.get('1stbullet', '')
            replacements[f'{{{{experience[{i}].2ndbullet}}}}'] = exp.get('2ndbullet', '')
        
        # Handle project section
        projects = data.get('project', [])
        for i, proj in enumerate(projects):
            replacements[f'{{{{project[{i}].name}}}}'] = proj.get('name', '')
            replacements[f'{{{{project[{i}].dates}}}}'] = proj.get('dates', '')
            replacements[f'{{{{project[{i}].1stbullet}}}}'] = proj.get('1stbullet', '')
            replacements[f'{{{{project[{i}].2ndbullet}}}}'] = proj.get('2ndbullet', '')
            replacements[f'{{{{project[{i}].3dbullet}}}}'] = proj.get('3dbullet', '')
            replacements[f'{{{{project[{i}].4thbullet}}}}'] = proj.get('4thbullet', '')
            replacements[f'{{{{project[{i}].link}}}}'] = proj.get('link', '')
        
        # Handle education section
        education = data.get('education', [])
        for i, edu in enumerate(education):
            replacements[f'{{{{education[{i}].branch}}}}'] = edu.get('branch', '')
            replacements[f'{{{{education[{i}].cgpa}}}}'] = edu.get('cgpa', '')
        
        # Handle responsibilities section
        responsibilities = data.get('responsibilities', [])
        for i, resp in enumerate(responsibilities):
            name = resp.get('name', '')
            org = resp.get('organization', '')
            replacements[f'{{{{responsibilities[{i}].name}}}}'] = name
            replacements[f'{{{{responsibilities[{i}].organization}}}}'] = org
            replacements[f'{{{{responsibilities[{i}].bullet}}}}'] = resp.get('bullet', '')
        
        # Replace all placeholders
        replace_all_placeholders(doc, replacements)
        
        # Save the document
        doc.save(DOCX_FILENAME)
        logger.info(f'DOCX successfully saved at: {DOCX_FILENAME}')
        return DOCX_FILENAME

    except Exception as e:
        logger.error(f"Exception in generate_docx: {e}")
        raise

@app.route('/generate_docx', methods=['POST'])
def create_docx():
    try:
        # Read raw request body
        raw_data = request.data.decode('utf-8')
        logger.debug('Raw request data: %s', raw_data)

        # Strip markdown code fences if present
        raw_data = re.sub(r'^```(?:json)?\s*', '', raw_data)
        raw_data = re.sub(r'\s*```$', '', raw_data)
        logger.debug('Cleaned request data: %s', raw_data)

        # Parse JSON
        data = json.loads(raw_data)
        if not data:
            logger.warning('No JSON data after cleaning')
            return jsonify({"error": "No data provided"}), 400
        print(data)
        # Generate DOCX from cleaned JSON
        generate_docx(data)
        download_url = request.host_url.rstrip('/') + '/download_resume'
        return jsonify({
            "message": "DOCX generated successfully!",
            "download_url": download_url
        }), 200

    except Exception as e:
        logger.exception("Error in /generate_docx")
        return jsonify({"error": "DOCX generation failed", "details": str(e)}), 500

@app.route('/download_resume', methods=['GET'])
def download_resume():
    if os.path.exists(DOCX_FILENAME):
        logger.info('Serving generated resume DOCX')
        return send_file(DOCX_FILENAME, 
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name='resume.docx')
    else:
        logger.warning('DOCX not found for download')
        return jsonify({"error": "DOCX not found"}), 404

@app.route('/logs', methods=['GET'])
def view_logs():
    try:
        log_path = "app.log"
        if not os.path.exists(log_path):
            return jsonify({"error": "Log file not found"}), 404

        with open(log_path, 'r') as log_file:
            lines = log_file.readlines()[-100:]
            return jsonify({"log": ''.join(lines)}), 200

    except Exception as e:
        logger.error(f"Failed to retrieve logs: {e}")
        return jsonify({"error": "Could not retrieve logs"}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
